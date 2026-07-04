"""Embedded-image extraction from ``.docx`` — the modern ``<w:drawing>``
and legacy ``<w:pict>`` picture forms → :class:`ImageAlt` placeholders,
with the raster bytes preserved on :attr:`DocumentIR.assets`
(ARCHITECTURE.md, the I1 step).

The modern path is exercised through python-docx's real
``add_picture`` so the test proves genuine relationship resolution and
media-part bytes end to end — not a synthetic drawing that a green test
could pass while the real ``r:embed`` → ``word/media`` path is broken
(the lesson of the MTEF real-fixture regression).  The VML path and the
no-blip / unresolvable-bytes edge cases are hand-built because Word, not
python-docx, is what emits those shapes.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest

pytest.importorskip("docx")
pytest.importorskip("lxml")

from docx import Document  # noqa: E402
from lxml import etree  # noqa: E402

from brailix.input.docx import parse_docx  # noqa: E402
from brailix.ir.document import ImageAlt, Paragraph  # noqa: E402

_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_V_NS = "urn:schemas-microsoft-com:vml"
_O_NS = "urn:schemas-microsoft-com:office:office"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

# Opaque bytes for the paths that only store + retrieve the blob (VML,
# unresolvable-blip): the content is never parsed there, only round-tripped
# through DocumentIR.assets, so any recognisable marker works.
_FAKE_IMG = b"\x89PNG\r\n\x1a\n-not-a-real-png-just-a-marker-"


@pytest.fixture
def real_png() -> bytes:
    """A well-formed PNG python-docx's own image parser will accept — its
    ``add_picture`` reads the header for dimensions, so the modern-drawing
    tests need a genuine raster, not a marker blob. Built with Pillow (the
    ``graphics`` extra); the modern tests skip when it isn't installed."""
    pytest.importorskip("PIL")
    from PIL import Image

    buf = BytesIO()
    Image.new("RGB", (4, 4), (200, 30, 30)).save(buf, format="PNG")
    return buf.getvalue()


def _parse(doc: Document, tmp_path: Path):
    out = tmp_path / "doc.docx"
    doc.save(str(out))
    return parse_docx(out, language="zh-CN", profile="cn_current")


def _first_docpr(drawing_run) -> object:
    """The ``<wp:docPr>`` element inside a python-docx picture run."""
    for elem in drawing_run._r.iter():
        if etree.QName(elem).localname == "docPr":
            return elem
    raise AssertionError("no docPr in picture run")


def _relate_image(doc: Document, blob: bytes, ext: str = "png") -> str:
    """Add ``blob`` as a ``word/media`` image part + an IMAGE relationship;
    return its rId (mirrors the OLE helper in ``test_docx``)."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.packuri import PackURI
    from docx.opc.part import Part

    existing = [
        n
        for n in doc.part.package.iter_parts()
        if "/word/media/image" in str(n.partname)
    ]
    idx = len(existing) + 1
    part = Part(
        partname=PackURI(f"/word/media/image{idx}.{ext}"),
        content_type=f"image/{ext}",
        blob=blob,
        package=doc.part.package,
    )
    return doc.part.relate_to(part, RT.IMAGE)


class TestModernDrawing:
    """``<w:drawing>`` (the DrawingML picture Word writes today)."""

    def test_picture_becomes_image_alt_with_asset(
        self, tmp_path: Path, real_png: bytes
    ) -> None:
        doc = Document()
        doc.add_paragraph("前")
        doc.add_picture(BytesIO(real_png))
        doc.add_paragraph("后")

        ir = _parse(doc, tmp_path)

        images = [b for b in ir.blocks if isinstance(b, ImageAlt)]
        assert len(images) == 1
        assert images[0].target == "media/image1.png"
        # The bytes rode out on DocumentIR.assets under the same name the
        # placeholder references, so nothing external has to be re-read.
        assert ir.assets["media/image1.png"] == real_png
        # Prose around the picture is preserved and ordered.
        prose = [b.text for b in ir.blocks if isinstance(b, Paragraph)]
        assert prose == ["前", "后"]

    def test_docpr_descr_is_the_alt_text(
        self, tmp_path: Path, real_png: bytes
    ) -> None:
        doc = Document()
        run = doc.add_paragraph().add_run()
        run.add_picture(BytesIO(real_png))
        _first_docpr(run).set("descr", "一张地图")

        ir = _parse(doc, tmp_path)

        (image,) = [b for b in ir.blocks if isinstance(b, ImageAlt)]
        assert image.text == "一张地图"
        assert image.target == "media/image1.png"

    def test_blank_alt_falls_back_to_media_stem(
        self, tmp_path: Path, real_png: bytes
    ) -> None:
        # add_picture writes docPr name="Picture 1" but no descr/title; we
        # clear name too so the *stem* fallback is what's under test.
        doc = Document()
        run = doc.add_paragraph().add_run()
        run.add_picture(BytesIO(real_png))
        docpr = _first_docpr(run)
        docpr.set("name", "")

        ir = _parse(doc, tmp_path)

        (image,) = [b for b in ir.blocks if isinstance(b, ImageAlt)]
        assert image.text == "image1"

    def test_picture_splits_a_paragraph(
        self, tmp_path: Path, real_png: bytes
    ) -> None:
        # A picture inline with text: braille has no inline images, so the
        # placeholder surfaces as its own block between the text slices.
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("左")
        para.add_run().add_picture(BytesIO(real_png))
        para.add_run("右")

        ir = _parse(doc, tmp_path)

        kinds = [type(b).__name__ for b in ir.blocks]
        assert kinds == ["Paragraph", "ImageAlt", "Paragraph"]
        assert ir.blocks[0].text == "左"
        assert ir.blocks[2].text == "右"


class TestLegacyVmlPict:
    """``<w:pict>`` — the VML picture Word still writes for some documents
    (and LibreOffice conversions emit)."""

    def test_imagedata_becomes_image_alt(self, tmp_path: Path) -> None:
        doc = Document()
        rid = _relate_image(doc, _FAKE_IMG)
        para = doc.add_paragraph()
        pict_xml = (
            f'<w:r xmlns:w="{_W_NS}" xmlns:r="{_R_NS}" '
            f'xmlns:v="{_V_NS}" xmlns:o="{_O_NS}">'
            f"<w:pict>"
            f'<v:shape id="s1" alt="红点" type="#_x0000_t75">'
            f'<v:imagedata r:id="{rid}" o:title="dot"/>'
            f"</v:shape>"
            f"</w:pict>"
            f"</w:r>"
        )
        para._p.append(etree.fromstring(pict_xml))

        ir = _parse(doc, tmp_path)

        (image,) = [b for b in ir.blocks if isinstance(b, ImageAlt)]
        assert image.text == "红点"
        assert image.target == "media/image1.png"
        assert ir.assets["media/image1.png"] == _FAKE_IMG


class TestEdgeCases:
    def test_drawing_without_blip_is_skipped(self, tmp_path: Path) -> None:
        # A drawing that references no raster (a shape / chart / text box) has
        # nothing to preserve as a tactile image — it stays out of scope, and
        # the surrounding paragraph is untouched.
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("图形说明")
        drawing_xml = (
            f'<w:r xmlns:w="{_W_NS}" xmlns:wp="{_WP_NS}">'
            f"<w:drawing><wp:inline>"
            f'<wp:docPr id="1" name="Shape"/>'
            f"</wp:inline></w:drawing>"
            f"</w:r>"
        )
        para._p.append(etree.fromstring(drawing_xml))

        ir = _parse(doc, tmp_path)

        assert not [b for b in ir.blocks if isinstance(b, ImageAlt)]
        assert [b.text for b in ir.blocks if isinstance(b, Paragraph)] == [
            "图形说明"
        ]

    def test_unresolvable_blip_keeps_placeholder_without_target(
        self, tmp_path: Path
    ) -> None:
        # The picture's relationship is dangling (bytes unrecoverable), but its
        # existence must still survive: an alt-text placeholder with no target.
        doc = Document()
        para = doc.add_paragraph()
        drawing_xml = (
            f'<w:r xmlns:w="{_W_NS}" xmlns:wp="{_WP_NS}" '
            f'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
            f'xmlns:r="{_R_NS}">'
            f"<w:drawing><wp:inline>"
            f'<wp:docPr id="1" name="Picture 1" descr="丢失的图"/>'
            f"<a:graphic><a:graphicData>"
            f'<a:blip r:embed="rIdMissing"/>'
            f"</a:graphicData></a:graphic>"
            f"</wp:inline></w:drawing>"
            f"</w:r>"
        )
        para._p.append(etree.fromstring(drawing_xml))

        ir = _parse(doc, tmp_path)

        (image,) = [b for b in ir.blocks if isinstance(b, ImageAlt)]
        assert image.text == "丢失的图"
        assert image.target is None
        assert ir.assets == {}
